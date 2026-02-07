# app.py - Flask Web Application for Placeholder Pipeline
# Provides web UI for document processing and placeholder filling

import os
import io
import json
import tempfile
import datetime
from flask import Flask, request, jsonify, render_template, send_file
from flask_cors import CORS
from werkzeug.utils import secure_filename

# Import pipeline modules
from modules.databricks_parser import DatabricksParser
from modules.smart_placeholder_filler import SmartPlaceholderFiller
from modules.pdf_extractor import PDFExtractor
from modules.databricks_client import upload_to_volume, get_databricks_client

app = Flask(__name__)
CORS(app)

# Configuration
DATABRICKS_VOLUME_PATH = "/Volumes/workspace/default/test/"

# Use a fixed upload folder that survives server restarts
UPLOAD_FOLDER = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'uploads')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

ALLOWED_EXTENSIONS = {'docx', 'pdf'}

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB max


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


# ============== PAGE ROUTES ==============

@app.route('/test')
def test():
    """Simple test endpoint"""
    import sys
    # Write to file since stdout might be buffered
    with open('debug.txt', 'a') as f:
        f.write(f"TEST ENDPOINT HIT at {datetime.datetime.now()}\n")
    print("!!! TEST ENDPOINT HIT !!!", flush=True)
    sys.stdout.flush()
    return "Flask is working! Check debug.txt for logs."

@app.route('/')
def index():
    """Main page with step-by-step wizard"""
    return render_template('index.html')


# ============== DATABRICKS API ==============

@app.route('/api/databricks/files', methods=['GET'])
def list_databricks_files():
    """List files in Databricks volume"""
    try:
        parser = DatabricksParser()
        
        # Try to list files using Databricks SQL
        # Note: This requires proper Databricks connection
        try:
            from databricks import sql
            
            if all([parser.host, parser.http_path, parser.access_token]):
                with sql.connect(
                    server_hostname=parser.host,
                    http_path=parser.http_path,
                    access_token=parser.access_token
                ) as connection:
                    with connection.cursor() as cursor:
                        # List files in the volume
                        query = f"LIST '{DATABRICKS_VOLUME_PATH}'"
                        cursor.execute(query)
                        results = cursor.fetchall()
                        
                        files = []
                        for row in results:
                            # row typically contains: path, name, size, modificationTime
                            files.append({
                                'name': row.name if hasattr(row, 'name') else str(row[0]),
                                'path': row.path if hasattr(row, 'path') else f"{DATABRICKS_VOLUME_PATH}{row[0]}",
                                'size': row.size if hasattr(row, 'size') else 0
                            })
                        
                        return jsonify({
                            'success': True,
                            'files': files,
                            'path': DATABRICKS_VOLUME_PATH
                        })
        except ImportError:
            pass
        except Exception as e:
            app.logger.warning(f"Databricks connection failed: {e}")
        
        # Fallback: Return mock data for demo
        return jsonify({
            'success': True,
            'files': [
                {'name': 'Report.docx', 'path': f'{DATABRICKS_VOLUME_PATH}Report.docx', 'size': 713317},
                {'name': 'Annual Update.docx', 'path': f'{DATABRICKS_VOLUME_PATH}Annual Update.docx', 'size': 2637121},
            ],
            'path': DATABRICKS_VOLUME_PATH,
            'mock': True
        })
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/databricks/upload', methods=['POST'])
def upload_to_databricks():
    """Upload file to Databricks volume"""
    if 'file' not in request.files:
        return jsonify({'success': False, 'error': 'No file provided'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'success': False, 'error': 'No file selected'}), 400
    
    if not allowed_file(file.filename):
        return jsonify({'success': False, 'error': 'File type not allowed'}), 400
    
    try:
        filename = secure_filename(file.filename)
        
        # Save locally first
        local_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(local_path)
        
        # Upload to Databricks volume using SDK
        try:
            databricks_path = f"{DATABRICKS_VOLUME_PATH}{filename}"
            uploaded_path = upload_to_volume(local_path, databricks_path)
            
            return jsonify({
                'success': True,
                'filename': filename,
                'local_path': local_path,
                'databricks_path': uploaded_path
            })
        except Exception as e:
            app.logger.error(f"Databricks upload failed: {e}")
            # Fallback to local path if upload fails, but warn user
            return jsonify({
                'success': True,
                'filename': filename,
                'local_path': local_path,
                'databricks_path': None,
                'warning': f"Failed to upload to Databricks: {str(e)}. Using local file."
            })
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


# ============== PROCESSING API ==============

@app.route('/api/process/create-placeholders', methods=['POST'])
def create_placeholders():
    """Process DOCX and create placeholders"""
    data = request.json
    file_path = data.get('file_path')
    local_file = data.get('local_file')
    
    if not file_path and not local_file:
        return jsonify({'success': False, 'error': 'No file specified'}), 400
    
    try:
        parser = DatabricksParser()
        
        # Try to upload to Databricks first if it's a local file
        path_to_process = file_path or local_file
        
        if local_file and os.path.exists(local_file):
            filename = os.path.basename(local_file)
            target_path = f"{DATABRICKS_VOLUME_PATH}{filename}"
            try:
                # Upload to Databricks
                path_to_process = upload_to_volume(local_file, target_path)
                app.logger.info(f"Uploaded {local_file} to {path_to_process}")
            except Exception as e:
                app.logger.warning(f"Failed to verify/upload to Databricks: {e}")
                # Continue with local path - DatabricksParser will handle fallback or fail
        
        # Parse and classify
        result = parser.parse_and_classify(path_to_process)
        
        if not result.classification_result:
            return jsonify({
                'success': False,
                'error': 'Failed to classify document'
            }), 500
        
        # Get placeholders from classification
        placeholders = []
        seen = set()
        for r in result.classification_result.replacements:
            # Extract placeholder name from {{NAME}}
            import re
            match = re.search(r'\{\{([A-Z0-9_]+)\}\}', r.placeholder)
            if match and match.group(1) not in seen:
                seen.add(match.group(1))
                placeholders.append({
                    'name': match.group(1),
                    'placeholder': r.placeholder,
                    'category': r.category,
                    'original_preview': r.original[:100] + '...' if len(r.original) > 100 else r.original
                })
        
        # Rebuild document with placeholders
        template_bytes = None
        if result.document_bytes and result.classification_result.replacements:
            from modules.document_rebuilder import DocumentRebuilder
            
            rebuilder = DocumentRebuilder()
            build_result = rebuilder.rebuild_from_bytes(
                result.document_bytes,
                result.classification_result.replacements
            )
            
            if build_result.success:
                template_bytes = build_result.document_bytes
                
                # Save template for later
                timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                template_filename = f"template_{timestamp}.docx"
                template_path = os.path.join(app.config['UPLOAD_FOLDER'], template_filename)
                
                with open(template_path, 'wb') as f:
                    f.write(template_bytes)
                
                # Extract ACTUAL placeholders from the built template
                from docx import Document as DocxDocument
                import re
                doc = DocxDocument(template_path)
                pattern = re.compile(r'\{\{([A-Z0-9_]+)\}\}')
                actual_placeholders = set()
                
                for para in doc.paragraphs:
                    matches = pattern.findall(para.text)
                    actual_placeholders.update(matches)
                
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                matches = pattern.findall(para.text)
                                actual_placeholders.update(matches)
                
                # Build placeholder list from actual content
                placeholders = [
                    {'name': p, 'placeholder': f'{{{{{p}}}}}', 'category': p}
                    for p in sorted(actual_placeholders)
                ]
                
                return jsonify({
                    'success': True,
                    'placeholders': placeholders,
                    'template_file': template_filename,
                    'template_path': template_path,
                    'replacements_applied': build_result.replacements_applied,
                    'replacements_failed': build_result.replacements_failed
                })
        
        return jsonify({
            'success': False,
            'error': 'Failed to create template',
            'placeholders': placeholders
        }), 500
        
    except Exception as e:
        app.logger.error(f"Error creating placeholders: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/process/fill-placeholders', methods=['POST'])
def fill_placeholders():
    """Fill selected placeholders with PDF content"""
    import sys
    
    def log(msg):
        """Log to both stdout and debug.txt"""
        print(msg, flush=True)
        with open('debug.txt', 'a') as f:
            f.write(f"{msg}\n")
    
    log("\n" + "="*60)
    log("FILL-PLACEHOLDERS ENDPOINT HIT")
    log("="*60)
    
    try:
        template_file = request.form.get('template_file')
        placeholder_mappings = json.loads(request.form.get('mappings', '{}'))
        
        log(f"Template file: {template_file}")
        log(f"Placeholder mappings from frontend: {placeholder_mappings}")
        log(f"Request files: {list(request.files.keys())}")
        
        if not template_file:
            log("ERROR: No template file specified")
            return jsonify({'success': False, 'error': 'No template file specified'}), 400
        
        template_path = os.path.join(app.config['UPLOAD_FOLDER'], template_file)
        print(f"Template path: {template_path}", flush=True)
        print(f"Template exists: {os.path.exists(template_path)}", flush=True)
        
        if not os.path.exists(template_path):
            print("ERROR: Template file not found", flush=True)
            return jsonify({'success': False, 'error': 'Template file not found'}), 404
        
        # Save uploaded PDFs and build mapping
        pdf_mapping = {}
        for placeholder_name in placeholder_mappings:
            pdf_key = f'pdf_{placeholder_name}'
            print(f"Looking for PDF key: {pdf_key}", flush=True)
            if pdf_key in request.files:
                pdf_file = request.files[pdf_key]
                print(f"  Found file: {pdf_file.filename}", flush=True)
                if pdf_file and allowed_file(pdf_file.filename):
                    pdf_filename = secure_filename(pdf_file.filename)
                    pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], pdf_filename)
                    pdf_file.save(pdf_path)
                    pdf_mapping[placeholder_name] = pdf_path
                    print(f"  Saved PDF for {placeholder_name}: {pdf_filename}", flush=True)
            else:
                print(f"  NOT FOUND in request.files", flush=True)
        
        print(f"Final PDF mapping: {pdf_mapping}", flush=True)
        sys.stdout.flush()
        
        if not pdf_mapping:
            print("ERROR: No PDF files uploaded", flush=True)
            return jsonify({'success': False, 'error': 'No PDF files uploaded'}), 400
        
        # Use smart filler
        print("Creating SmartPlaceholderFiller...", flush=True)
        filler = SmartPlaceholderFiller()
        
        # Generate output filename
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"filled_{timestamp}.docx"
        output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
        
        print(f"Calling filler.fill_with_pdf...", flush=True)
        print(f"  template_path: {template_path}", flush=True)
        print(f"  pdf_mapping: {pdf_mapping}", flush=True)
        print(f"  output_path: {output_path}", flush=True)
        sys.stdout.flush()
        
        result = filler.fill_with_pdf(template_path, pdf_mapping, output_path)
        
        print(f"Smart filler result: found={result.placeholders_found}, filled={result.placeholders_filled}", flush=True)
        print(f"Errors: {result.errors}", flush=True)
        print(f"Details: {result.details}", flush=True)
        sys.stdout.flush()
        
        if result.success:
            return jsonify({
                'success': True,
                'output_file': output_filename,
                'placeholders_found': result.placeholders_found,
                'placeholders_filled': result.placeholders_filled,
                'details': result.details,
                'errors': result.errors
            })
        else:
            return jsonify({
                'success': False,
                'error': 'Failed to fill placeholders',
                'errors': result.errors
            }), 500
            
    except Exception as e:
        app.logger.error(f"Error filling placeholders: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/upload/template', methods=['POST'])
def upload_template():
    """Upload existing template with placeholders"""
    if 'file' not in request.files:
        return jsonify({'success': False, 'error': 'No file provided'}), 400
    
    file = request.files['file']
    if file.filename == '' or not allowed_file(file.filename):
        return jsonify({'success': False, 'error': 'Invalid file'}), 400
    
    try:
        filename = secure_filename(file.filename)
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        saved_filename = f"uploaded_{timestamp}_{filename}"
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], saved_filename)
        file.save(filepath)
        
        # Extract placeholders from the document
        from docx import Document
        import re
        
        doc = Document(filepath)
        placeholders = set()
        pattern = re.compile(r'\{\{([A-Z0-9_]+)\}\}')
        
        for para in doc.paragraphs:
            matches = pattern.findall(para.text)
            placeholders.update(matches)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        matches = pattern.findall(para.text)
                        placeholders.update(matches)
        
        return jsonify({
            'success': True,
            'template_file': saved_filename,
            'template_path': filepath,
            'placeholders': [{'name': p, 'placeholder': f'{{{{{p}}}}}'}for p in sorted(placeholders)]
        })
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/download/<filename>')
def download_file(filename):
    """Download processed document"""
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], secure_filename(filename))
    
    if not os.path.exists(filepath):
        return jsonify({'success': False, 'error': 'File not found'}), 404
    
    return send_file(
        filepath,
        as_attachment=True,
        download_name=filename
    )


if __name__ == '__main__':
    print("=" * 50)
    print("Placeholder Pipeline Web UI")
    print("=" * 50)
    print(f"Upload folder: {UPLOAD_FOLDER}")
    print(f"Databricks path: {DATABRICKS_VOLUME_PATH}")
    print("Starting server at http://localhost:5000")
    print("=" * 50)
    
    app.run(debug=True, host='0.0.0.0', port=5000)
