# databricks_parser.py - Databricks ai_parse_document Integration
# Extracts text and structure using Databricks AI functions

import logging
import os
import json
from typing import Dict, Any, Optional
from dataclasses import dataclass

try:
    from databricks import sql
except ImportError:
    sql = None

import sys
sys.path.insert(0, os.path.dirname(os.path.dirname(__file__)))
from config import CLASSIFICATION_PROMPT_TEMPLATE
from modules.content_classifier import ContentClassifier, ClassificationResult

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@dataclass
class DatabricksParserResult:
    """Result from Databricks parsing"""
    text: str
    markdown: str
    metadata: Dict[str, Any]
    classification_result: Optional[ClassificationResult] = None
    document_bytes: Optional[bytes] = None

class DatabricksParser:
    """
    Parses documents using Databricks ai_parse_document and applies 
    existing classification logic.
    """

    def __init__(self, host: str = None, http_path: str = None, access_token: str = None):
        """
        Initialize the Databricks parser.
        
        Args:
            host: Databricks host (from env if None)
            http_path: Databricks HTTP path (from env if None)
            access_token: Databricks access token (from env if None)
        """
        self.host = host or os.getenv("DATABRICKS_HOST")
        self.http_path = http_path or os.getenv("DATABRICKS_HTTP_PATH")
        self.access_token = access_token or os.getenv("DATABRICKS_TOKEN")
        
        if not all([self.host, self.http_path, self.access_token]):
            logger.warning("Databricks credentials not fully set. Connection will fail unless mocked.")

    def parse_and_classify(self, file_url: str) -> DatabricksParserResult:
        """
        Parse a document using ai_parse_document and classify its content.
        """
        # 1. Extract content using Databricks
        extracted_data = self._extract_content_from_databricks(file_url)
        
        # Extracted data contains 'content' (bytes) and 'parsed' (dict)
        document_bytes = extracted_data.get('content')
        parsed_dict = extracted_data.get('parsed', {})
        
        logger.info(f"Extracted keys: {list(parsed_dict.keys())}")
        
        # 2. Classify content using existing logic
        classifier = ContentClassifier()
        # process the markdown text as it likely preserves structure better
        # Keys observed: ['document', 'error_status', 'metadata']
        markdown_text = parsed_dict.get('markdown')
        
        doc_obj = parsed_dict.get('document')
        if not markdown_text and doc_obj:
            if isinstance(doc_obj, dict):
                logger.info(f"Document object keys: {list(doc_obj.keys())}")
                # Try common keys
                markdown_text = doc_obj.get('markdown') or doc_obj.get('content') or doc_obj.get('text')
                
                # Fallback: Construct from elements if available
                if not markdown_text:
                    elements = doc_obj.get('elements') or doc_obj.get('chunks')
                    if elements and isinstance(elements, list):
                        logger.info(f"Reconstructing text from {len(elements)} elements")
                        # Sort by page/id if needed, but assuming list order is reading order
                        text_parts = []
                        for el in elements:
                            if isinstance(el, dict) and 'content' in el:
                                content = el['content']
                                if content is not None:
                                    text_parts.append(str(content))
                        
                        if text_parts:
                            markdown_text = "\n\n".join(text_parts)
            else:
                markdown_text = doc_obj
                
        if not markdown_text:
             markdown_text = parsed_dict.get('text') or ''
        
        # Ensure it is a string
        if not isinstance(markdown_text, str):
            # If it's still a dict/obj, likely we failed to find the text field
            # Taking str() of a huge dict is bad for classification
            if isinstance(markdown_text, (dict, list)):
                logger.warning("Markdown text is a structure, extracting raw string representation (suboptimal)")
            markdown_text = str(markdown_text)
            
        # DEBUG: Save extracted text to file
        with open("debug_extracted_text.txt", "w", encoding="utf-8") as f:
            f.write(markdown_text)
        logger.info(f"Saved extracted text to debug_extracted_text.txt ({len(markdown_text)} chars)")
            
        classification = classifier.classify_document(markdown_text)
        
        return DatabricksParserResult(
            text=markdown_text,
            markdown=markdown_text,
            metadata=parsed_dict.get('metadata', {}),
            classification_result=classification,
            document_bytes=document_bytes
        )

    def _extract_content_from_databricks(self, file_url: str) -> Dict[str, Any]:
        """
        Call execute ai_parse_document on Databricks.
        """
        # Determine if we can run real Databricks calls
        can_run_real = (sql is not None) and all([self.host, self.http_path, self.access_token])
        
        if not can_run_real:
            reason = "databricks-sql-connector not installed" if sql is None else "credentials missing"
            logger.info(f"Cannot run real Databricks query ({reason}). Using MOCK mode.")
            return self._mock_response(file_url)

        try:
            with sql.connect(
                server_hostname=self.host,
                http_path=self.http_path,
                access_token=self.access_token
            ) as connection:
                with connection.cursor() as cursor:
                    # Construct the query
                    # ai_parse_document requires BINARY input. 
                    # We use read_files to read the content from the Volume/path as binary.
                    
                    # Escape the path just in case
                    safe_path = file_url.replace("'", "''")
                    
                    # Query to get BOTH original content and parsed result
                    query = f"""
                        SELECT content, ai_parse_document(content) as parsed 
                        FROM read_files('{safe_path}', format => 'binaryFile')
                        LIMIT 1
                    """
                    logger.info(f"Executing Query: {query}")
                    
                    cursor.execute(query)
                    result = cursor.fetchone()
                    
                    if result:
                        # result is a Row object, fields accessed by name or index
                        file_content = result.content
                        parsed_data = result.parsed
                        
                        if parsed_data:
                            # Parse the returned struct/JSON
                            # The return type depends on Databricks version
                            
                            # If parsed is a string (JSON), decode it
                            if isinstance(parsed_data, str):
                                try:
                                    parsed_data = json.loads(parsed_data)
                                except json.JSONDecodeError:
                                    logger.error("Failed to decode JSON response from Databricks")
                                    raise
                            
                            return {
                                "content": file_content,
                                "parsed": parsed_data
                            }
                        else:
                             raise ValueError("ai_parse_document returned None")
                    else:
                        raise ValueError("No result returned from query")

        except Exception as e:
            logger.error(f"Databricks execution failed: {e}")
            logger.info("Falling back to MOCK response for testing.")
            return self._mock_response(file_url)

    def _mock_response(self, file_url: str) -> Dict[str, Any]:
        """
        Return a mock response simulating ai_parse_document output.
        If file_url is a local path that exists, try to read it using python-docx.
        """
        # Check if it's a local file
        if os.path.exists(file_url):
            try:
                from docx import Document
                import io
                
                # Read bytes
                with open(file_url, 'rb') as f:
                    content_bytes = f.read()
                
                # Extract text using python-docx
                doc = Document(file_url)
                text_parts = []
                for para in doc.paragraphs:
                    text_parts.append(para.text)
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                text_parts.append(para.text)
                                
                full_text = "\n\n".join(text_parts)
                
                logger.info(f"Fallback: Read local file {file_url} ({len(full_text)} chars)")
                
                return {
                    "content": content_bytes,
                    "parsed": {
                        "text": full_text,
                        "markdown": full_text,  # Simple text as markdown
                        "metadata": {"source": file_url, "pages": 1, "mock_mode": True, "local_fallback": True}
                    }
                }
            except Exception as e:
                logger.warning(f"Failed to read local file {file_url} in fallback mode: {e}")
        
        # Default mock response
        return {
            "content": None,  # Cannot provide valid DOCX bytes in mock mode
            "parsed": {
                "text": "This is a mock parsed text. \n\nWe recommend you diversify your portfolio.",
                "markdown": "# Mock Document\n\nThis is a mock parsed text.\n\nWe recommend you diversify your portfolio.\n\nYour goal is to save $1 million.",
                "metadata": {"source": file_url, "pages": 1, "mock_mode": True}
            }
        }

if __name__ == "__main__":
    # Test script integration
    import argparse
    import datetime
    
    parser = argparse.ArgumentParser(description="Databricks Document Parser with Placeholder Filling")
    parser.add_argument("--file", help="Path to file (URL/Volume path)", default="dbfs:/tmp/test.docx")
    parser.add_argument("--pdf-risk", help="Path to local PDF file for {{RISK}} placeholder content")
    parser.add_argument("--output", help="Output filename (default: auto-generated with timestamp)")
    args = parser.parse_args()
    
    print(f"Testing Databricks Parser with file: {args.file}")
    
    db_parser = DatabricksParser()
    result = db_parser.parse_and_classify(args.file)
    
    print("\n=== Extracted Markdown ===")
    print(result.markdown)
    
    print("\n=== Classification Result ===")
    if result.classification_result:
        for r in result.classification_result.replacements:
            print(f"- [Original]: {r.original}")
            print(f"  [Replacement]: {r.placeholder}")
            print(f"  [Category]: {r.category}")
            print(f"  [Confidence]: {r.confidence}")
            
    print(f"\nDEBUG: Has document_bytes: {bool(result.document_bytes)}")
    if result.document_bytes:
        print(f"DEBUG: document_bytes length: {len(result.document_bytes)}")
    print(f"DEBUG: Has classification_result: {bool(result.classification_result)}")
    if result.classification_result:
        print(f"DEBUG: replacements count: {len(result.classification_result.replacements)}")
    
    # Check if in mock mode
    is_mock_mode = result.metadata.get("mock_mode", False)
    if is_mock_mode:
        print("\n=== MOCK MODE ===")
        print("Document bytes not available in mock mode - cannot rebuild document.")
        print("To test document rebuilding, provide a valid Databricks volume path.")
            
    # Rebuild document if content is available
    template_docx_bytes = None
    if result.document_bytes and result.classification_result and result.classification_result.replacements:
        print("\n=== Rebuilding Document (Creating Placeholders) ===")
        from modules.document_rebuilder import DocumentRebuilder
        
        rebuilder = DocumentRebuilder()
        build_result = rebuilder.rebuild_from_bytes(
            result.document_bytes,
            result.classification_result.replacements
        )
        
        print(f"Placeholders applied: {build_result.replacements_applied}")
        print(f"Placeholders failed: {build_result.replacements_failed}")
        
        if build_result.failed_items:
            print("\n=== FAILED REPLACEMENTS ===")
            for item in build_result.failed_items:
                # Truncate strictly for display
                display_item = item[:80] + "..." if len(item) > 80 else item
                print(f"- {display_item}")
        
        if build_result.success:
            template_docx_bytes = build_result.document_bytes
        else:
            print(f"FAILED to rebuild document: {build_result.errors}")
    
    # === PHASE 2: Fill placeholders with PDF content (SMART - using GPT) ===
    if args.pdf_risk and template_docx_bytes:
        print("\n=== Smart Filling Placeholders with PDF Content (GPT) ===")
        
        try:
            from modules.smart_placeholder_filler import SmartPlaceholderFiller
            
            # Build PDF mapping
            pdf_mapping = {"RISK_STATEMENT": args.pdf_risk}
            
            print(f"Using GPT to extract relevant content from: {args.pdf_risk}")
            print("GPT will analyze context around {{RISK_STATEMENT}} placeholder and extract matching content...")
            
            # Save template to temp file for smart filler
            import tempfile
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                tmp.write(template_docx_bytes)
                temp_path = tmp.name
            
            try:
                # Use smart filler with GPT
                smart_filler = SmartPlaceholderFiller()
                fill_result = smart_filler.fill_with_pdf(
                    temp_path,
                    pdf_mapping,
                    output_path=None  # Get bytes back
                )
                
                print(f"\nPlaceholders found: {fill_result.placeholders_found}")
                print(f"Placeholders filled: {fill_result.placeholders_filled}")
                
                if fill_result.details:
                    print("\n=== Extracted Content Summaries ===")
                    for name, summary in fill_result.details.items():
                        print(f"  {{{{{name}}}}}: {summary}")
                
                if fill_result.errors:
                    print(f"\nWarnings/Errors: {fill_result.errors}")
                
                if fill_result.success and fill_result.document_bytes:
                    # Save the final filled document
                    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                    output_filename = args.output or f"smart_filled_{timestamp}.docx"
                    
                    try:
                        with open(output_filename, "wb") as f:
                            f.write(fill_result.document_bytes)
                        print(f"\nSUCCESS: Smart-filled document saved to: {os.path.abspath(output_filename)}")
                    except IOError as e:
                        print(f"ERROR: Failed to save file: {e}")
                else:
                    print(f"FAILED to fill placeholders: {fill_result.errors}")
            finally:
                # Clean up temp file
                os.unlink(temp_path)
                
        except ImportError as e:
            print(f"ERROR: Missing dependency: {e}")
            print("Install required packages: pip install pdfplumber openai")
    
    elif template_docx_bytes and not args.pdf_risk:
        # Save template without filling (original behavior)
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = args.output or f"template_output_{timestamp}.docx"
        
        try:
            with open(output_filename, "wb") as f:
                f.write(template_docx_bytes)
            print(f"\nSUCCESS: Template document saved to: {os.path.abspath(output_filename)}")
            print("TIP: Use --pdf-risk to fill {{RISK}} placeholder with PDF content")
        except IOError as e:
            print(f"ERROR: Failed to save file: {e}")
