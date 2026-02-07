# document_parser.py - DOCX Document Parsing Module
# Extracts text, tables, and structure from Word documents

import logging
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass, field
from docx import Document
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


@dataclass
class TextElement:
    """Represents a text element with its location and content"""
    text: str
    element_type: str  # 'paragraph', 'table_cell', 'header', 'footer'
    paragraph_index: int = -1
    table_index: int = -1
    row_index: int = -1
    cell_index: int = -1
    run_indices: List[int] = field(default_factory=list)
    style_name: Optional[str] = None
    is_bold: bool = False
    is_italic: bool = False
    font_size: Optional[float] = None
    
    def get_path(self) -> str:
        """Returns a unique path identifier for this element"""
        if self.element_type == 'paragraph':
            return f"p[{self.paragraph_index}]"
        elif self.element_type == 'table_cell':
            return f"table[{self.table_index}]/row[{self.row_index}]/cell[{self.cell_index}]"
        elif self.element_type in ('header', 'footer'):
            return f"{self.element_type}/p[{self.paragraph_index}]"
        return f"unknown[{self.paragraph_index}]"


@dataclass
class DocumentStructure:
    """Represents the complete structure of a parsed document"""
    text_elements: List[TextElement] = field(default_factory=list)
    full_text: str = ""
    table_count: int = 0
    paragraph_count: int = 0
    image_count: int = 0
    has_headers: bool = False
    has_footers: bool = False
    metadata: Dict[str, Any] = field(default_factory=dict)
    
    def get_text_for_analysis(self) -> str:
        """Returns combined text for LLM analysis"""
        return self.full_text
    
    def get_elements_containing(self, text: str) -> List[TextElement]:
        """Find all elements containing the given text"""
        return [elem for elem in self.text_elements if text in elem.text]


class DocumentParser:
    """
    Parses DOCX files and extracts text with structure preservation.
    Adapted from LayIE-LLM methodology for layout-aware extraction.
    """
    
    def __init__(self):
        self.current_doc: Optional[Document] = None
        self.structure: Optional[DocumentStructure] = None
    
    def parse(self, docx_path: str) -> DocumentStructure:
        """
        Parse a DOCX file and extract its structure.
        
        Args:
            docx_path: Path to the DOCX file
            
        Returns:
            DocumentStructure containing all extracted elements
        """
        logger.info(f"Parsing document: {docx_path}")
        
        self.current_doc = Document(docx_path)
        self.structure = DocumentStructure()
        
        # Extract paragraphs
        self._extract_paragraphs()
        
        # Extract tables
        self._extract_tables()
        
        # Extract headers and footers
        self._extract_headers_footers()
        
        # Count images
        self._count_images()
        
        # Build full text
        self._build_full_text()
        
        # Extract metadata
        self._extract_metadata()
        
        logger.info(f"Parsed {self.structure.paragraph_count} paragraphs, "
                   f"{self.structure.table_count} tables, "
                   f"{self.structure.image_count} images")
        
        return self.structure
    
    def _extract_paragraphs(self) -> None:
        """Extract all paragraphs from the document body"""
        for idx, paragraph in enumerate(self.current_doc.paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue
            
            # Determine formatting
            is_bold = any(run.bold for run in paragraph.runs if run.bold is not None)
            is_italic = any(run.italic for run in paragraph.runs if run.italic is not None)
            
            # Get style
            style_name = paragraph.style.name if paragraph.style else None
            
            # Get run indices for this paragraph
            run_indices = list(range(len(paragraph.runs)))
            
            element = TextElement(
                text=text,
                element_type='paragraph',
                paragraph_index=idx,
                run_indices=run_indices,
                style_name=style_name,
                is_bold=is_bold,
                is_italic=is_italic
            )
            
            self.structure.text_elements.append(element)
            self.structure.paragraph_count += 1
    
    def _extract_tables(self) -> None:
        """Extract text from all tables"""
        for table_idx, table in enumerate(self.current_doc.tables):
            self.structure.table_count += 1
            
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    if not cell_text:
                        continue
                    
                    # Check formatting of first paragraph in cell
                    is_bold = False
                    is_italic = False
                    if cell.paragraphs:
                        first_para = cell.paragraphs[0]
                        is_bold = any(run.bold for run in first_para.runs if run.bold is not None)
                        is_italic = any(run.italic for run in first_para.runs if run.italic is not None)
                    
                    element = TextElement(
                        text=cell_text,
                        element_type='table_cell',
                        table_index=table_idx,
                        row_index=row_idx,
                        cell_index=cell_idx,
                        is_bold=is_bold,
                        is_italic=is_italic
                    )
                    
                    self.structure.text_elements.append(element)
    
    def _extract_headers_footers(self) -> None:
        """Extract text from headers and footers"""
        # Process all sections
        for section in self.current_doc.sections:
            # Headers
            header = section.header
            if header and header.paragraphs:
                self.structure.has_headers = True
                for idx, para in enumerate(header.paragraphs):
                    text = para.text.strip()
                    if text:
                        element = TextElement(
                            text=text,
                            element_type='header',
                            paragraph_index=idx
                        )
                        self.structure.text_elements.append(element)
            
            # Footers
            footer = section.footer
            if footer and footer.paragraphs:
                self.structure.has_footers = True
                for idx, para in enumerate(footer.paragraphs):
                    text = para.text.strip()
                    if text:
                        element = TextElement(
                            text=text,
                            element_type='footer',
                            paragraph_index=idx
                        )
                        self.structure.text_elements.append(element)
    
    def _count_images(self) -> None:
        """Count images in the document"""
        # Count inline shapes (images)
        image_count = 0
        for rel in self.current_doc.part.rels.values():
            if "image" in rel.target_ref:
                image_count += 1
        
        self.structure.image_count = image_count
    
    def _build_full_text(self) -> None:
        """Build the complete text representation for analysis"""
        text_parts = []
        
        # Add paragraph texts with context markers
        for elem in self.structure.text_elements:
            if elem.element_type == 'paragraph':
                # Add style context for headings
                if elem.style_name and 'Heading' in elem.style_name:
                    text_parts.append(f"\n[{elem.style_name}] {elem.text}\n")
                else:
                    text_parts.append(elem.text)
            elif elem.element_type == 'table_cell':
                text_parts.append(f"[Table{elem.table_index + 1}] {elem.text}")
            elif elem.element_type in ('header', 'footer'):
                text_parts.append(f"[{elem.element_type.capitalize()}] {elem.text}")
        
        self.structure.full_text = '\n'.join(text_parts)
    
    def _extract_metadata(self) -> None:
        """Extract document metadata"""
        core_props = self.current_doc.core_properties
        
        self.structure.metadata = {
            'title': core_props.title or '',
            'author': core_props.author or '',
            'created': str(core_props.created) if core_props.created else '',
            'modified': str(core_props.modified) if core_props.modified else '',
            'subject': core_props.subject or '',
            'keywords': core_props.keywords or ''
        }
    
    def get_element_by_path(self, path: str) -> Optional[TextElement]:
        """Find an element by its path identifier"""
        for elem in self.structure.text_elements:
            if elem.get_path() == path:
                return elem
        return None
    
    def find_elements_with_text(self, search_text: str, exact: bool = False) -> List[TextElement]:
        """
        Find all elements containing the search text.
        
        Args:
            search_text: Text to search for
            exact: If True, match exact text; if False, partial match
            
        Returns:
            List of matching TextElement objects
        """
        results = []
        search_lower = search_text.lower()
        
        for elem in self.structure.text_elements:
            if exact:
                if elem.text == search_text:
                    results.append(elem)
            else:
                if search_lower in elem.text.lower():
                    results.append(elem)
        
        return results


def parse_document(docx_path: str) -> DocumentStructure:
    """
    Convenience function to parse a document.
    
    Args:
        docx_path: Path to the DOCX file
        
    Returns:
        DocumentStructure containing parsed content
    """
    parser = DocumentParser()
    return parser.parse(docx_path)
