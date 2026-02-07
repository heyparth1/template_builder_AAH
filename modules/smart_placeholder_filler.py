# smart_placeholder_filler.py - Context-Aware Placeholder Filling with LLM
# Uses GPT to extract relevant content from PDF based on document context

import logging
import re
import io
import json
from typing import Dict, List, Optional, Tuple
from dataclasses import dataclass, field
from docx import Document
from docx.text.paragraph import Paragraph

import sys
import os
sys.path.insert(0, os.path.dirname(os.path.dirname(__file__)))
from modules.model_interaction import GPTClient, create_client
from modules.pdf_extractor import PDFExtractor

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


@dataclass
class PlaceholderContext:
    """Context information for a placeholder"""
    placeholder_name: str  # e.g., "RISK"
    before_text: str       # Text before the placeholder
    after_text: str        # Text after the placeholder
    section_heading: str   # Nearest heading if available
    full_paragraph: str    # The paragraph containing the placeholder


@dataclass
class SmartFillerResult:
    """Result of smart placeholder filling"""
    success: bool
    placeholders_found: int
    placeholders_filled: int
    output_path: Optional[str] = None
    document_bytes: Optional[bytes] = None
    details: Dict[str, str] = field(default_factory=dict)  # placeholder -> extracted content summary
    errors: List[str] = field(default_factory=list)


CONTEXT_EXTRACTION_PROMPT = """You are an expert at extracting relevant content from documents based on context.

Given the following:
1. PLACEHOLDER CONTEXT: The surrounding text in the target document where content needs to be inserted
2. SOURCE CONTENT: The full content from a PDF file

Your task: Extract ONLY the portion of the SOURCE CONTENT that is most relevant to fill the placeholder based on the context.

## PLACEHOLDER CONTEXT
Placeholder Name: {placeholder_name}
Section/Heading: {section_heading}
Text Before Placeholder: {before_text}
Text After Placeholder: {after_text}
Full Paragraph: {full_paragraph}

## SOURCE CONTENT (from PDF)
{pdf_content}

## INSTRUCTIONS
1. Analyze the context AND the placeholder name to understand what type of content is expected
2. Extract the most relevant portion from the SOURCE CONTENT that fits this context
3. Keep the extracted content coherent and complete (don't cut off mid-sentence)
4. If the source content has sections, extract the appropriate section
5. Preserve formatting like bullet points if relevant

Return your response as JSON:
{{
    "extracted_content": "The relevant content extracted from the PDF that should fill this placeholder",
    "reasoning": "Brief explanation of why this content was selected"
}}
"""


class SmartPlaceholderFiller:
    """
    Fills placeholders using GPT to extract contextually relevant content.
    Reads the context around each placeholder and extracts matching content from PDF.
    """
    
    PLACEHOLDER_PATTERN = re.compile(r'\{\{([A-Z0-9_]+)\}\}')
    CONTEXT_CHARS = 500  # Characters of context to extract before/after placeholder
    
    def __init__(self, gpt_client: GPTClient = None):
        """
        Initialize the smart filler.
        
        Args:
            gpt_client: Optional pre-configured GPT client
        """
        self.gpt_client = gpt_client or create_client()
        self.pdf_extractor = PDFExtractor()
    
    def fill_with_pdf(
        self,
        docx_path: str,
        pdf_mapping: Dict[str, str],  # {"RISK": "path/to/risk.pdf", "GOALS": "path/to/goals.pdf"}
        output_path: Optional[str] = None
    ) -> SmartFillerResult:
        """
        Fill placeholders in a DOCX using content extracted from PDFs.
        Uses GPT to match content based on context.
        
        Args:
            docx_path: Path to DOCX with placeholders
            pdf_mapping: Dict mapping placeholder names to PDF file paths
            output_path: Optional output path for filled document
            
        Returns:
            SmartFillerResult with status and output
        """
        logger.info(f"Smart filling placeholders in: {docx_path}")
        logger.info(f"PDF mappings: {list(pdf_mapping.keys())}")
        
        errors = []
        details = {}
        placeholders_found = 0
        placeholders_filled = 0
        
        try:
            doc = Document(docx_path)
            
            # Extract PDF contents upfront
            pdf_contents = {}
            for placeholder_name, pdf_path in pdf_mapping.items():
                logger.info(f"Extracting PDF for {placeholder_name}: {pdf_path}")
                result = self.pdf_extractor.extract(pdf_path)
                if result.success:
                    pdf_contents[placeholder_name] = result.text
                    logger.info(f"  Extracted {len(result.text)} chars from {result.page_count} pages")
                else:
                    errors.append(f"Failed to extract {pdf_path}: {result.error}")
            
            # Find all placeholders with their context
            placeholder_contexts = self._find_placeholders_with_context(doc)
            placeholders_found = len(placeholder_contexts)
            logger.info(f"Found {placeholders_found} placeholders with context")
            
            # Log all found placeholder names for debugging
            found_names = [ctx.placeholder_name for ctx in placeholder_contexts]
            logger.info(f"Placeholder names in document: {found_names}")
            logger.info(f"PDF content available for: {list(pdf_contents.keys())}")
            
            # Process each placeholder - with flexible matching
            content_mapping = {}
            for ctx in placeholder_contexts:
                # Direct match
                if ctx.placeholder_name in pdf_contents:
                    pdf_content = pdf_contents[ctx.placeholder_name]
                else:
                    # Try prefix matching (RISK_STATEMENT matches RISK_STATEMENT_1, RISK_STATEMENT_2, etc.)
                    pdf_content = None
                    for pdf_key in pdf_contents.keys():
                        if ctx.placeholder_name.startswith(pdf_key):
                            pdf_content = pdf_contents[pdf_key]
                            logger.info(f"Prefix match: {ctx.placeholder_name} -> {pdf_key}")
                            break
                
                if pdf_content:
                    logger.info(f"Processing placeholder: {{{{{ctx.placeholder_name}}}}}")
                    
                    # Use GPT to extract relevant content
                    extracted = self._extract_relevant_content(
                        ctx,
                        pdf_content
                    )
                    
                    if extracted:
                        content_mapping[ctx.placeholder_name] = extracted
                        details[ctx.placeholder_name] = extracted[:100] + "..." if len(extracted) > 100 else extracted
                        logger.info(f"  Extracted {len(extracted)} chars for {{{{{ctx.placeholder_name}}}}}")
                    else:
                        errors.append(f"Failed to extract content for {ctx.placeholder_name}")
            
            
            # Apply replacements
            for placeholder_name, content in content_mapping.items():
                if self._replace_placeholder(doc, placeholder_name, content):
                    placeholders_filled += 1
            
            # Save or return bytes
            if output_path:
                doc.save(output_path)
                logger.info(f"Saved filled document to: {output_path}")
                
                return SmartFillerResult(
                    success=True,
                    placeholders_found=placeholders_found,
                    placeholders_filled=placeholders_filled,
                    output_path=output_path,
                    details=details,
                    errors=errors
                )
            else:
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                
                return SmartFillerResult(
                    success=True,
                    placeholders_found=placeholders_found,
                    placeholders_filled=placeholders_filled,
                    document_bytes=buffer.getvalue(),
                    details=details,
                    errors=errors
                )
                
        except Exception as e:
            logger.error(f"Error in smart filling: {e}")
            return SmartFillerResult(
                success=False,
                placeholders_found=placeholders_found,
                placeholders_filled=placeholders_filled,
                details=details,
                errors=[str(e)]
            )
    
    def fill_from_bytes(
        self,
        docx_bytes: bytes,
        pdf_mapping: Dict[str, str]
    ) -> SmartFillerResult:
        """
        Fill placeholders in DOCX bytes using PDF content.
        
        Args:
            docx_bytes: DOCX file as bytes
            pdf_mapping: Dict mapping placeholder names to PDF paths
            
        Returns:
            SmartFillerResult with document bytes
        """
        logger.info("Smart filling from document bytes")
        
        errors = []
        details = {}
        placeholders_found = 0
        placeholders_filled = 0
        
        try:
            buffer = io.BytesIO(docx_bytes)
            doc = Document(buffer)
            
            # Extract PDF contents
            pdf_contents = {}
            for placeholder_name, pdf_path in pdf_mapping.items():
                result = self.pdf_extractor.extract(pdf_path)
                if result.success:
                    pdf_contents[placeholder_name] = result.text
                else:
                    errors.append(f"Failed to extract {pdf_path}: {result.error}")
            
            # Find placeholders with context
            placeholder_contexts = self._find_placeholders_with_context(doc)
            placeholders_found = len(placeholder_contexts)
            
            # Process each placeholder with GPT
            content_mapping = {}
            for ctx in placeholder_contexts:
                if ctx.placeholder_name in pdf_contents:
                    extracted = self._extract_relevant_content(
                        ctx,
                        pdf_contents[ctx.placeholder_name]
                    )
                    if extracted:
                        content_mapping[ctx.placeholder_name] = extracted
                        details[ctx.placeholder_name] = extracted[:100] + "..."
            
            # Apply replacements
            for placeholder_name, content in content_mapping.items():
                if self._replace_placeholder(doc, placeholder_name, content):
                    placeholders_filled += 1
            
            # Return as bytes
            output_buffer = io.BytesIO()
            doc.save(output_buffer)
            output_buffer.seek(0)
            
            return SmartFillerResult(
                success=True,
                placeholders_found=placeholders_found,
                placeholders_filled=placeholders_filled,
                document_bytes=output_buffer.getvalue(),
                details=details,
                errors=errors
            )
            
        except Exception as e:
            logger.error(f"Error in smart filling: {e}")
            return SmartFillerResult(
                success=False,
                placeholders_found=placeholders_found,
                placeholders_filled=placeholders_filled,
                details=details,
                errors=[str(e)]
            )
    
    def _find_placeholders_with_context(self, doc: Document) -> List[PlaceholderContext]:
        """Find all placeholders and extract their surrounding context."""
        contexts = []
        current_heading = "Document Start"
        all_paragraphs = []
        
        # Collect all paragraphs with their indices
        for i, para in enumerate(doc.paragraphs):
            all_paragraphs.append((i, para))
            
            # Track headings
            if para.style and 'Heading' in para.style.name:
                current_heading = para.text.strip()
        
        # Find placeholders and build context
        seen_placeholders = set()
        
        # 1. Check Paragraphs
        for i, para in all_paragraphs:
            matches = self.PLACEHOLDER_PATTERN.findall(para.text)
            for match in matches:
                if match in seen_placeholders:
                    continue
                seen_placeholders.add(match)
                
                # Get before context (previous paragraphs)
                before_parts = []
                for j in range(max(0, i-3), i):
                    before_parts.append(doc.paragraphs[j].text)
                before_text = "\n".join(before_parts)[-self.CONTEXT_CHARS:]
                
                # Get after context (next paragraphs)
                after_parts = []
                for j in range(i+1, min(len(doc.paragraphs), i+4)):
                    after_parts.append(doc.paragraphs[j].text)
                after_text = "\n".join(after_parts)[:self.CONTEXT_CHARS]
                
                # Find nearest heading
                section_heading = current_heading
                for j in range(i-1, -1, -1):
                    h_para = doc.paragraphs[j]
                    if h_para.style and 'Heading' in h_para.style.name:
                        section_heading = h_para.text.strip()
                        break
                
                contexts.append(PlaceholderContext(
                    placeholder_name=match,
                    before_text=before_text,
                    after_text=after_text,
                    section_heading=section_heading,
                    full_paragraph=para.text
                ))

        # 2. Check Tables
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para in cell.paragraphs:
                        matches = self.PLACEHOLDER_PATTERN.findall(para.text)
                        for match in matches:
                            if match in seen_placeholders:
                                continue
                            
                            seen_placeholders.add(match)
                            
                            # For table cells, context is the cell content and surrounding cells/rows
                            # We'll use the cell's text as full_paragraph, and try to get some before/after context
                            
                            # Before text: previous paragraphs in the cell, or previous cells/rows
                            before_text_parts = []
                            # Add text from previous paragraphs in the same cell
                            cell_paras = list(cell.paragraphs)
                            current_para_idx = cell_paras.index(para) if para in cell_paras else -1
                            if current_para_idx > 0:
                                before_text_parts.extend([p.text for p in cell_paras[:current_para_idx]])
                            
                            # Add text from previous cells in the same row (up to 2 cells)
                            for prev_cell_idx in range(max(0, cell_idx - 2), cell_idx):
                                if prev_cell_idx < cell_idx: # Ensure it's a previous cell
                                    before_text_parts.append(row.cells[prev_cell_idx].text)
                            
                            # Add text from previous row (last cell)
                            if row_idx > 0:
                                prev_row = table.rows[row_idx - 1]
                                if prev_row.cells:
                                    before_text_parts.append(prev_row.cells[-1].text)
                            
                            before_text = "\n".join(before_text_parts)[-self.CONTEXT_CHARS:]

                            # After text: next paragraphs in the cell, or next cells/rows
                            after_text_parts = []
                            # Add text from next paragraphs in the same cell
                            if current_para_idx != -1 and current_para_idx < len(cell_paras) - 1:
                                after_text_parts.extend([p.text for p in cell_paras[current_para_idx + 1:]])

                            # Add text from next cells in the same row (up to 2 cells)
                            for next_cell_idx in range(cell_idx + 1, min(len(row.cells), cell_idx + 3)):
                                after_text_parts.append(row.cells[next_cell_idx].text)
                            
                            # Add text from next row (first cell)
                            if row_idx < len(table.rows) - 1:
                                next_row = table.rows[row_idx + 1]
                                if next_row.cells:
                                    after_text_parts.append(next_row.cells[0].text)

                            after_text = "\n".join(after_text_parts)[:self.CONTEXT_CHARS]
                            
                            # Section heading for tables can be the last heading before the table
                            table_section_heading = current_heading # Default to the last seen heading
                            # More precise: find heading immediately preceding the table
                            # This would require iterating through document elements, which is more complex.
                            # For now, we'll use the last known heading.
                            
                            contexts.append(PlaceholderContext(
                                placeholder_name=match,
                                before_text=before_text,
                                after_text=after_text,
                                section_heading=table_section_heading,
                                full_paragraph=para.text # The paragraph containing the placeholder
                            ))
        
        return contexts
    
    def _extract_relevant_content(self, context: PlaceholderContext, pdf_content: str) -> Optional[str]:
        """Use GPT to extract relevant content from PDF based on context."""
        prompt = CONTEXT_EXTRACTION_PROMPT.format(
            placeholder_name=context.placeholder_name,
            section_heading=context.section_heading,
            before_text=context.before_text,
            after_text=context.after_text,
            full_paragraph=context.full_paragraph,
            pdf_content=pdf_content[:15000]  # Limit to avoid token limits
        )
        
        try:
            response = self.gpt_client.call_api(
                prompt=prompt,
                system_message="You are an expert at extracting and matching content from documents. Return only valid JSON.",
                max_tokens=8192,
                temperature=0.2,
                json_mode=True
            )
            
            if response.success:
                # Try to parse JSON response
                extracted = self._parse_gpt_response(response.content)
                
                if extracted:
                    # Clean up the extracted content for better formatting
                    extracted = self._normalize_text(extracted)
                    return extracted
                else:
                    logger.error("Failed to parse GPT response")
                    return None
            else:
                logger.error(f"GPT call failed: {response.error}")
                return None
                
        except Exception as e:
            logger.error(f"Error extracting content: {e}")
            return None
    
    def _parse_gpt_response(self, content: str) -> Optional[str]:
        """
        Parse GPT response with fallback strategies for malformed JSON.
        """
        # Strategy 1: Try direct JSON parse
        try:
            result = json.loads(content)
            if isinstance(result, list):
                # If it's a list (maybe list of strings or objects), try to join them or get first
                if result and isinstance(result[0], str):
                    return "\n\n".join(result)
                elif result and isinstance(result[0], dict):
                    return result[0].get("extracted_content", "")
                return str(result)
            return result.get("extracted_content", "")
        except json.JSONDecodeError:
            pass
        except AttributeError:
            pass
        
        # Strategy 2: Try to extract content using regex
        # Look for "extracted_content": "..." pattern
        try:
            # Match content between "extracted_content": " and the next unescaped quote
            pattern = r'"extracted_content"\s*:\s*"((?:[^"\\]|\\.)*)\"'
            match = re.search(pattern, content, re.DOTALL)
            if match:
                extracted = match.group(1)
                # Unescape common escape sequences
                extracted = extracted.replace('\\"', '"')
                extracted = extracted.replace('\\n', '\n')
                extracted = extracted.replace('\\r', '\r')
                extracted = extracted.replace('\\t', '\t')
                logger.info("Extracted content using regex fallback")
                return extracted
        except Exception as e:
            logger.warning(f"Regex extraction failed: {e}")
        
        # Strategy 3: Look for content between first { and last }
        try:
            start = content.find('"extracted_content"')
            if start >= 0:
                # Find the colon and opening quote
                colon_pos = content.find(':', start)
                if colon_pos >= 0:
                    quote_start = content.find('"', colon_pos + 1)
                    if quote_start >= 0:
                        # Find the closing - look for ", or "}
                        remaining = content[quote_start + 1:]
                        # Simple approach: find "reasoning" or end of object
                        end_markers = ['", "reasoning"', '","reasoning"', '"}']
                        min_end = len(remaining)
                        for marker in end_markers:
                            pos = remaining.find(marker)
                            if pos >= 0 and pos < min_end:
                                min_end = pos
                        
                        if min_end < len(remaining):
                            extracted = remaining[:min_end]
                            extracted = extracted.replace('\\"', '"')
                            logger.info("Extracted content using string search fallback")
                            return extracted
        except Exception as e:
            logger.warning(f"String search extraction failed: {e}")
        
        logger.error(f"All parsing strategies failed. Raw content preview: {content[:500]}...")
        return None
    
    def _normalize_text(self, text: str) -> str:
        """
        Normalize text to pretty print in Word.
        - Preserves paragraph usage (double newlines)
        - Cleans up excessive internal whitespace
        """
        if not text:
            return text
        
        # Replace carriage returns with simple newlines
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # Ensure max 2 newlines (paragraph break)
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # We DO NOT replace newlines with spaces anymore to preserve layout!
        # text = text.replace('\n', ' ') <-- REMOVED
        
        # Strip leading/trailing whitespace
        text = text.strip()
        
        return text
    
    def _replace_placeholder(self, doc: Document, placeholder: str, content: str) -> bool:
        """Replace a placeholder throughout the document."""
        placeholder_text = f"{{{{{placeholder}}}}}"
        replaced = False
        
        # Replace in paragraphs
        for para in doc.paragraphs:
            if placeholder_text in para.text:
                self._replace_in_paragraph(para, placeholder_text, content)
                replaced = True

        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if placeholder_text in para.text:
                            self._replace_in_paragraph(para, placeholder_text, content)
                            replaced = True
        
        return replaced
    
    def _replace_in_paragraph(self, para: Paragraph, placeholder: str, content: str) -> None:
        """Replace placeholder in a paragraph."""
        # Try run-by-run first
        for run in para.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, content)
                return
        
        # If spans runs, reconstruct
        if placeholder in para.text:
            new_text = para.text.replace(placeholder, content)
            if para.runs:
                para.runs[0].text = new_text
                for run in para.runs[1:]:
                    run.text = ""


if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description="Smart Placeholder Filler with GPT")
    parser.add_argument("--template", required=True, help="Path to DOCX template with placeholders")
    parser.add_argument("--pdf-risk", help="Path to PDF for {{RISK}} placeholder")
    parser.add_argument("--output", help="Output path for filled document")
    args = parser.parse_args()
    
    pdf_mapping = {}
    if args.pdf_risk:
        pdf_mapping["RISK_STATEMENT"] = args.pdf_risk
    
    if not pdf_mapping:
        print("Error: At least one --pdf-* argument required")
        exit(1)
    
    filler = SmartPlaceholderFiller()
    result = filler.fill_with_pdf(args.template, pdf_mapping, args.output)
    
    print(f"\nPlaceholders found: {result.placeholders_found}")
    print(f"Placeholders filled: {result.placeholders_filled}")
    
    if result.details:
        print("\n=== Extracted Content Summaries ===")
        for name, summary in result.details.items():
            print(f"{{{{{name}}}}}: {summary}")
    
    if result.errors:
        print(f"\nErrors: {result.errors}")
    
    if result.output_path:
        print(f"\nOutput saved to: {result.output_path}")
