# document_rebuilder.py - Document Reconstruction Module
# Applies placeholder replacements while preserving formatting

import logging
import re
import copy
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass
from docx import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.table import Table, _Cell
from docx.oxml.ns import qn
from lxml import etree
import io

import sys
import os
sys.path.insert(0, os.path.dirname(os.path.dirname(__file__)))
from modules.content_classifier import Replacement

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


@dataclass
class ReplacementResult:
    """Result of document rebuilding"""
    success: bool
    replacements_applied: int
    replacements_failed: int
    output_path: Optional[str] = None
    document_bytes: Optional[bytes] = None
    failed_items: List[str] = None
    errors: List[str] = None
    
    def __post_init__(self):
        if self.errors is None:
            self.errors = []
        if self.failed_items is None:
            self.failed_items = []


class DocumentRebuilder:
    """
    Rebuilds DOCX documents with placeholder replacements.
    Preserves all formatting, tables, images, headers, and footers.
    """
    
    def __init__(self):
        self.replacements_applied = 0
        self.replacements_failed = 0
        self.errors = []
    
    def rebuild(
        self,
        input_path: str,
        replacements: List[Replacement],
        output_path: Optional[str] = None
    ) -> ReplacementResult:
        """
        Rebuild document with placeholders.
        
        Args:
            input_path: Path to input DOCX file
            replacements: List of replacements to apply
            output_path: Optional path for output file
            
        Returns:
            ReplacementResult with status and output
        """
        logger.info(f"Rebuilding document: {input_path}")
        logger.info(f"Applying {len(replacements)} replacements")
        
        self.replacements_applied = 0
        self.replacements_failed = 0
        self.errors = []
        failed_items = []
        
        try:
            # Load document
            doc = Document(input_path)
            
            # Sort replacements by length (longest first)
            sorted_replacements = sorted(
                replacements,
                key=lambda r: len(r.original),
                reverse=True
            )
            
            # Apply replacements to different parts of the document
            for replacement in sorted_replacements:
                applied = False
                
                # Try paragraphs
                if self._replace_in_paragraphs(doc, replacement):
                    applied = True
                
                # Try tables
                if self._replace_in_tables(doc, replacement):
                    applied = True
                
                # Try headers and footers
                if self._replace_in_headers_footers(doc, replacement):
                    applied = True
                
                if applied:
                    self.replacements_applied += 1
                    logger.debug(f"Applied: {replacement.original[:30]}... -> {replacement.placeholder}")
                else:
                    self.replacements_failed += 1
                    failed_items.append(replacement.original)
                    logger.debug(f"Not found: {replacement.original[:30]}...")
            
            # Save or return bytes
            if output_path:
                doc.save(output_path)
                logger.info(f"Saved to: {output_path}")
                
                return ReplacementResult(
                    success=True,
                    replacements_applied=self.replacements_applied,
                    replacements_failed=self.replacements_failed,
                    output_path=output_path,
                    failed_items=failed_items,
                    errors=self.errors
                )
            else:
                # Return as bytes
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                
                return ReplacementResult(
                    success=True,
                    replacements_applied=self.replacements_applied,
                    replacements_failed=self.replacements_failed,
                    document_bytes=buffer.getvalue(),
                    failed_items=failed_items,
                    errors=self.errors
                )
                
        except Exception as e:
            logger.error(f"Error rebuilding document: {e}")
            return ReplacementResult(
                success=False,
                replacements_applied=self.replacements_applied,
                replacements_failed=self.replacements_failed,
                failed_items=failed_items,
                errors=[str(e)]
            )
    
    def _replace_in_paragraphs(self, doc: Document, replacement: Replacement) -> bool:
        """
        Replace text in document paragraphs.
        Handles text split across multiple runs.
        """
        applied = False
        
        # Try single paragraph replacement first (most common)
        for paragraph in doc.paragraphs:
            if self._replace_in_paragraph(paragraph, replacement):
                applied = True
                
        # If not applied, try multi-paragraph replacement
        if not applied and self._replace_multi_paragraph_in_list(doc.paragraphs, replacement):
            applied = True
            
        return applied

    def _replace_multi_paragraph_in_list(self, paragraphs: List[Paragraph], replacement: Replacement) -> bool:
        """
        Try to find and replace text that spans multiple paragraphs.
        """
        # 1. Inspect replacement.original for newlines
        parts = [p.strip() for p in replacement.original.splitlines() if p.strip()]
        if len(parts) <= 1:
            return False
            
        n_paras = len(paragraphs)
        n_parts = len(parts)
        
        # Iterate through paragraphs to find the sequence
        for i in range(n_paras - n_parts + 1):
            match = True
            
            for j in range(n_parts):
                doc_text = " ".join(paragraphs[i+j].text.split())
                part_text = " ".join(parts[j].split())
                
                # Check for containment (fuzzy match)
                if part_text not in doc_text:
                    match = False
                    break
            
            if match:
                logger.info("Found multi-paragraph match!")
                # Replace first paragraph
                # Note: This is a destructive replace, losing run formatting in first paragraph
                # But necessary for multi-paragraph blocks
                paragraphs[i].text = replacement.placeholder
                
                # Clear content of subsequent paragraphs
                # We leave the paragraphs themselves to preserve layout structure/spacing
                # but empty their content
                for k in range(1, n_parts):
                    paragraphs[i+k].text = ""
                    
                return True
                
        return False
    
    def _replace_in_paragraph(self, paragraph: Paragraph, replacement: Replacement) -> bool:
        """
        Replace text in a single paragraph while preserving formatting.
        """
        full_text = paragraph.text
        
        # Quick check using loose matching
        parts = replacement.original.split()
        if not parts:
            return False
            
        # If the first few words aren't in the text, skip (optimization)
        # Note: This might be too aggressive if formatting breaks words, but generally safe
        if parts[0] not in full_text and len(parts[0]) > 5:
             # Try stricter check before failing? 
             # Let's rely on _replace_across_runs for the heavy lifting
             pass

        # Strategy 1: Try simple run-by-run replacement first (Exact Match Only)
        # We keep this for simple cases where exact match exists in a single run, avoiding regex overhead
        for run in paragraph.runs:
            if replacement.original in run.text:
                run.text = run.text.replace(replacement.original, replacement.placeholder)
                return True
        
        # Strategy 2: Text spans multiple runs OR exact match failed due to whitespace
        # _replace_across_runs method now handles regex/fuzzy matching
        if self._replace_across_runs(paragraph, replacement):
            return True
            
        # Strategy 3: Normalized matching (ignore punctuation/whitespace)
        # Handles "long-term" vs "long term", "Client's" vs "Clients", etc.
        if self._replace_normalized(paragraph, replacement):
            logger.info(f"Normalized match applied for: {replacement.original[:30]}...")
            return True

        # Strategy 4: Fuzzy matching (slowest, use as last resort)
        # Handles smart quotes, soft hyphens, minor OCR quirks
        if self._replace_fuzzy(paragraph, replacement):
            logger.info(f"Fuzzy match applied for: {replacement.original[:30]}...")
            return True
            
        return False
    
    def _replace_normalized(self, paragraph: Paragraph, replacement: Replacement) -> bool:
        """
        Replace text by ignoring punctuation and whitespace.
        Useful for differences like:
        - "long-term" vs "long term"
        - "client's" vs "clients"
        - "hello, world" vs "hello world"
        """
        import re
        
        full_text = paragraph.text
        
        # Helper to normalize and build index map
        def normalize_with_map(text):
            normalized = []
            idx_map = [] # normalized_idx -> original_idx
            
            for i, char in enumerate(text):
                if char.isalnum():
                    normalized.append(char.lower())
                    idx_map.append(i)
            
            return "".join(normalized), idx_map

        # Normalize both texts
        doc_norm, doc_map = normalize_with_map(full_text)
        search_norm, _ = normalize_with_map(replacement.original)
        
        if not search_norm:
            return False
            
        # Check if search term exists in normalized doc text
        if search_norm not in doc_norm:
            return False
            
        # Find start index
        start_idx_norm = doc_norm.find(search_norm)
        end_idx_norm = start_idx_norm + len(search_norm)
        
        # Map back to original indices
        # start_idx_norm maps to the first alphanumeric character of the match
        # end_idx_norm - 1 maps to the last alphanumeric character of the match
        
        if start_idx_norm < len(doc_map) and (end_idx_norm - 1) < len(doc_map):
            start_pos = doc_map[start_idx_norm]
            # end_pos should be the index AFTER the last character
            end_pos = doc_map[end_idx_norm - 1] + 1
            
            # verify we aren't cutting off logically connected parts? 
            # (e.g. if original had punctuation at the end that we skipped)
            # logic: we replace from the first matched alphanum to the last matched alphanum
            # any punctuation *immediately* surrounding it might remain if it wasn't in the search string
            # but if the search string didn't have it, maybe we should keep it?
            # actually, if the replacement is a "correction", we might want to be careful.
            # But usually we are replacing a sentence/clause.
            
            return self._apply_replacement_at_range(paragraph, replacement, start_pos, end_pos)
            
        return False
    
    def _replace_across_runs(self, paragraph: Paragraph, replacement: Replacement) -> bool:
        """
        Handle replacement when text spans multiple runs.
        Preserves formatting of the first run.
        Uses whitespace-insensitive matching.
        """
        runs = paragraph.runs
        if not runs:
            return False
        
        # Build text with run boundaries
        run_texts = [(run, run.text) for run in runs]
        full_text = ''.join(text for _, text in run_texts)
        
        # Create a whitespace-insensitive pattern
        # Escape the original string, then replace escaped spaces with \s+ pattern
        # This handles space vs newline vs multiple spaces differences
        import re
        escaped = re.escape(replacement.original)
        # We want to match existing spaces in the search string to any whitespace sequence
        pattern_str = escaped.replace(r"\ ", r"\s+")
        
        # Should we match newlines in the original string too?
        # If original has \n, re.escape makes it \\n. 
        # simpler approach: split by whitespace and join with \s+
        parts = replacement.original.split()
        if not parts:
            return False
            
        # Construct pattern: part1\s+part2\s+part3 ...
        pattern_str = r"\s*".join(map(re.escape, parts))
        
        match = re.search(pattern_str, full_text, flags=re.MULTILINE | re.DOTALL)
        
        if not match:
            return False
            
        start_pos = match.start()
        end_pos = match.end()
        
        # Find which runs are affected
        current_pos = 0
        affected_runs = []
        
        for i, (run, text) in enumerate(run_texts):
            run_start = current_pos
            run_end = current_pos + len(text)
            
            # Check if this run overlaps with the replacement range
            # Overlap exists if start < run_end AND end > run_start
            if start_pos < run_end and end_pos > run_start:
                affected_runs.append((i, run, run_start, run_end))
            
            current_pos = run_end
        
        if not affected_runs:
            return False
        
        # Apply replacement
        try:
            if len(affected_runs) == 1:
                # Simple case: replacement within single run
                idx, run, run_start, run_end = affected_runs[0]
                
                # Careful with indices. run.text slice needs local indices.
                # The match might start before this run (shouldn't happen if logic correct) 
                # or end after (shouldn't happen).
                # Wait, if match spans multiple runs, len(affected_runs) > 1.
                # If len == 1, the whole match is inside this run.
                
                local_start = max(0, start_pos - run_start)
                local_end = min(len(run.text), end_pos - run_start)
                
                run.text = run.text[:local_start] + replacement.placeholder + run.text[local_end:]
            else:
                # Complex case: spans multiple runs
                # Put replacement in first affected run, clear others
                first_idx, first_run, first_start, first_end = affected_runs[0]
                
                local_start = max(0, start_pos - first_start)
                
                # Set first run text: keep prefix + placeholder
                first_run.text = first_run.text[:local_start] + replacement.placeholder
                
                # Middle runs - clear entirely
                for i in range(1, len(affected_runs) - 1):
                    _, run, _, _ = affected_runs[i]
                    run.text = ""
                    
                # Last affected run - keep trailing text
                last_idx, last_run, last_start, last_end = affected_runs[-1]
                local_end_in_last = end_pos - last_start
                # Ensure we don't slice negatively
                local_end_in_last = max(0, local_end_in_last)
                
                last_run.text = last_run.text[local_end_in_last:]
            
            return True
            
        except Exception as e:
            logger.warning(f"Error in cross-run replacement: {e}")
            return False

    def _replace_fuzzy(self, paragraph: Paragraph, replacement: Replacement) -> bool:
        """
        Attempt to find text using fuzzy matching (difflib).
        Useful for smart quotes, hidden variations, etc.
        """
        from difflib import SequenceMatcher
        
        text = paragraph.text
        # Threshold for matching (0.8 = 80% similarity)
        # We need a high threshold to avoid false positives
        # Lowered to 0.7 to handle encoding issues (e.g. £ vs ú)
        threshold = 0.7
        
        # We assume the replacement text is a substring. 
        # difflib finds longest common block.
        matcher = SequenceMatcher(None, text, replacement.original)
        match = matcher.find_longest_match(0, len(text), 0, len(replacement.original))
        
        # We want the match to vary effectively covers the whole original string
        # i.e. length of match is close to length of original
        if match.size / len(replacement.original) > threshold:
            # We found a strong partial match.
            # However, DocumentRebuilder needs exact indices to manipulate runs.
            # If the text in the doc is "slightly different", we can't just slice easily 
            # without determining the *actual* start/end in the doc text.
            
            # Since we found a match block, let's treat the *document range* as the target
            start_pos = match.a
            end_pos = match.a + match.size
            
            # Now apply replacement to this range (similar to _replace_across_runs logic)
            # We assume the "matched part" is what we want to replace.
            
            # Reuse logic from _replace_across_runs but with explicit start/end
            return self._apply_replacement_at_range(paragraph, replacement, start_pos, end_pos)
            
        return False

    def _apply_replacement_at_range(self, paragraph: Paragraph, replacement: Replacement, start_pos: int, end_pos: int) -> bool:
        """
        Helper to apply replacement at specific character indices in a paragraph.
        """
        runs = paragraph.runs
        run_texts = [(run, run.text) for run in runs]
        
        current_pos = 0
        affected_runs = []
        
        for i, (run, text) in enumerate(run_texts):
            run_start = current_pos
            run_end = current_pos + len(text)
            
            if start_pos < run_end and end_pos > run_start:
                affected_runs.append((i, run, run_start, run_end))
            
            current_pos = run_end
            
        if not affected_runs:
            return False
            
        try:
            if len(affected_runs) == 1:
                idx, run, run_start, run_end = affected_runs[0]
                local_start = max(0, start_pos - run_start)
                local_end = min(len(run.text), end_pos - run_start)
                run.text = run.text[:local_start] + replacement.placeholder + run.text[local_end:]
            else:
                first_idx, first_run, first_start, first_end = affected_runs[0]
                local_start = max(0, start_pos - first_start)
                first_run.text = first_run.text[:local_start] + replacement.placeholder
                
                for i in range(1, len(affected_runs) - 1):
                    _, run, _, _ = affected_runs[i]
                    run.text = ""
                    
                last_idx, last_run, last_start, last_end = affected_runs[-1]
                local_end_in_last = end_pos - last_start
                # Ensure we don't slice negatively
                local_end_in_last = max(0, local_end_in_last)
                
                last_run.text = last_run.text[local_end_in_last:]
            return True
        except Exception:
            return False
    
    def _replace_in_tables(self, doc: Document, replacement: Replacement) -> bool:
        """
        Replace text in all tables.
        """
        applied = False
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Try per-paragraph
                    cell_applied = False
                    for paragraph in cell.paragraphs:
                        if self._replace_in_paragraph(paragraph, replacement):
                            cell_applied = True
                            applied = True
                    
                    # Try multi-paragraph in this cell if not found yet
                    if not cell_applied and self._replace_multi_paragraph_in_list(cell.paragraphs, replacement):
                        applied = True
        
        return applied
    
    def _replace_in_headers_footers(self, doc: Document, replacement: Replacement) -> bool:
        """
        Replace text in headers and footers.
        """
        applied = False
        
        for section in doc.sections:
            # Header
            header = section.header
            if header:
                header_applied = False
                for paragraph in header.paragraphs:
                    if self._replace_in_paragraph(paragraph, replacement):
                        header_applied = True
                        applied = True
                        
                if not header_applied and self._replace_multi_paragraph_in_list(header.paragraphs, replacement):
                    applied = True
                
                # Tables in header
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            # Try per-paragraph
                            cell_applied = False
                            for paragraph in cell.paragraphs:
                                if self._replace_in_paragraph(paragraph, replacement):
                                    cell_applied = True
                                    applied = True
                            
                            if not cell_applied and self._replace_multi_paragraph_in_list(cell.paragraphs, replacement):
                                applied = True
            
            # Footer
            footer = section.footer
            if footer:
                footer_applied = False
                for paragraph in footer.paragraphs:
                    if self._replace_in_paragraph(paragraph, replacement):
                        footer_applied = True
                        applied = True
                
                if not footer_applied and self._replace_multi_paragraph_in_list(footer.paragraphs, replacement):
                    applied = True
                
                # Tables in footer
                for table in footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                if self._replace_in_paragraph(paragraph, replacement):
                                    applied = True
        
        return applied
    
    def rebuild_from_bytes(
        self,
        input_bytes: bytes,
        replacements: List[Replacement]
    ) -> ReplacementResult:
        """
        Rebuild document from bytes.
        
        Args:
            input_bytes: DOCX file as bytes
            replacements: List of replacements to apply
            
        Returns:
            ReplacementResult with document bytes
        """
        logger.info(f"Rebuilding document from bytes ({len(input_bytes)} bytes)")
        logger.info(f"Applying {len(replacements)} replacements")
        
        self.replacements_applied = 0
        self.replacements_failed = 0
        self.errors = []
        
        try:
            # Load document from bytes
            buffer = io.BytesIO(input_bytes)
            doc = Document(buffer)
            
            # Sort replacements by length (longest first)
            sorted_replacements = sorted(
                replacements,
                key=lambda r: len(r.original),
                reverse=True
            )
            
            # Apply replacements
            failed_items = []
            for replacement in sorted_replacements:
                applied = False
                
                if self._replace_in_paragraphs(doc, replacement):
                    applied = True
                if self._replace_in_tables(doc, replacement):
                    applied = True
                if self._replace_in_headers_footers(doc, replacement):
                    applied = True
                
                if applied:
                    self.replacements_applied += 1
                else:
                    self.replacements_failed += 1
                    failed_items.append(replacement.original)
            
            # Return as bytes
            output_buffer = io.BytesIO()
            doc.save(output_buffer)
            output_buffer.seek(0)
            
            return ReplacementResult(
                success=True,
                replacements_applied=self.replacements_applied,
                replacements_failed=self.replacements_failed,
                document_bytes=output_buffer.getvalue(),
                failed_items=failed_items,
                errors=self.errors
            )
            
        except Exception as e:
            logger.error(f"Error rebuilding document: {e}")
            return ReplacementResult(
                success=False,
                replacements_applied=self.replacements_applied,
                replacements_failed=self.replacements_failed,
                failed_items=failed_items if 'failed_items' in locals() else [],
                errors=[str(e)]
            )


def rebuild_document(
    input_path: str,
    replacements: List[Replacement],
    output_path: Optional[str] = None
) -> ReplacementResult:
    """
    Convenience function to rebuild a document.
    
    Args:
        input_path: Path to input DOCX
        replacements: List of replacements
        output_path: Optional output path
        
    Returns:
        ReplacementResult
    """
    rebuilder = DocumentRebuilder()
    return rebuilder.rebuild(input_path, replacements, output_path)
