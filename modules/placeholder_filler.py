# placeholder_filler.py - Fill Placeholders in DOCX Documents
# Replaces placeholders like {{RISK}} with content from PDF files

import logging
import re
import io
from typing import Dict, List, Optional, Tuple
from dataclasses import dataclass, field
from docx import Document
from docx.text.paragraph import Paragraph

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


@dataclass
class FillerResult:
    """Result of placeholder filling operation"""
    success: bool
    placeholders_found: int
    placeholders_filled: int
    output_path: Optional[str] = None
    document_bytes: Optional[bytes] = None
    unfilled: List[str] = field(default_factory=list)
    errors: List[str] = field(default_factory=list)


class PlaceholderFiller:
    """
    Fills placeholders in DOCX documents with provided content.
    Supports placeholders in format: {{PLACEHOLDER_NAME}}
    """
    
    # Regex to find placeholders like {{RISK}}, {{GOALS}}, etc.
    PLACEHOLDER_PATTERN = re.compile(r'\{\{([A-Z_]+)\}\}')
    
    def __init__(self):
        self.placeholders_found = 0
        self.placeholders_filled = 0
    
    def fill(
        self,
        docx_path: str,
        content_mapping: Dict[str, str],
        output_path: Optional[str] = None
    ) -> FillerResult:
        """
        Fill placeholders in a DOCX document.
        
        Args:
            docx_path: Path to DOCX template with placeholders
            content_mapping: Dict mapping placeholder names to content
                            e.g., {"RISK": "Risk assessment text..."}
            output_path: Optional path to save filled document
            
        Returns:
            FillerResult with status and output
        """
        logger.info(f"Filling placeholders in: {docx_path}")
        logger.info(f"Content mappings provided for: {list(content_mapping.keys())}")
        
        self.placeholders_found = 0
        self.placeholders_filled = 0
        unfilled = []
        errors = []
        
        try:
            doc = Document(docx_path)
            
            # Find all unique placeholders in the document
            all_placeholders = self._find_all_placeholders(doc)
            self.placeholders_found = len(all_placeholders)
            logger.info(f"Found {self.placeholders_found} unique placeholders: {all_placeholders}")
            
            # Fill each placeholder
            for placeholder in all_placeholders:
                if placeholder in content_mapping:
                    content = content_mapping[placeholder]
                    if self._replace_placeholder(doc, placeholder, content):
                        self.placeholders_filled += 1
                        logger.info(f"Filled placeholder: {{{{{placeholder}}}}}")
                    else:
                        unfilled.append(placeholder)
                        logger.warning(f"Failed to fill: {{{{{placeholder}}}}}")
                else:
                    unfilled.append(placeholder)
                    logger.warning(f"No content provided for: {{{{{placeholder}}}}}")
            
            # Save or return bytes
            if output_path:
                doc.save(output_path)
                logger.info(f"Saved filled document to: {output_path}")
                
                return FillerResult(
                    success=True,
                    placeholders_found=self.placeholders_found,
                    placeholders_filled=self.placeholders_filled,
                    output_path=output_path,
                    unfilled=unfilled,
                    errors=errors
                )
            else:
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                
                return FillerResult(
                    success=True,
                    placeholders_found=self.placeholders_found,
                    placeholders_filled=self.placeholders_filled,
                    document_bytes=buffer.getvalue(),
                    unfilled=unfilled,
                    errors=errors
                )
                
        except Exception as e:
            logger.error(f"Error filling placeholders: {e}")
            return FillerResult(
                success=False,
                placeholders_found=self.placeholders_found,
                placeholders_filled=self.placeholders_filled,
                unfilled=unfilled,
                errors=[str(e)]
            )
    
    def fill_from_bytes(
        self,
        docx_bytes: bytes,
        content_mapping: Dict[str, str]
    ) -> FillerResult:
        """
        Fill placeholders in a DOCX document from bytes.
        
        Args:
            docx_bytes: DOCX file as bytes
            content_mapping: Dict mapping placeholder names to content
            
        Returns:
            FillerResult with document bytes
        """
        logger.info("Filling placeholders from document bytes")
        
        self.placeholders_found = 0
        self.placeholders_filled = 0
        unfilled = []
        errors = []
        
        try:
            buffer = io.BytesIO(docx_bytes)
            doc = Document(buffer)
            
            # Find and fill placeholders
            all_placeholders = self._find_all_placeholders(doc)
            self.placeholders_found = len(all_placeholders)
            
            for placeholder in all_placeholders:
                if placeholder in content_mapping:
                    content = content_mapping[placeholder]
                    if self._replace_placeholder(doc, placeholder, content):
                        self.placeholders_filled += 1
                    else:
                        unfilled.append(placeholder)
                else:
                    unfilled.append(placeholder)
            
            # Return as bytes
            output_buffer = io.BytesIO()
            doc.save(output_buffer)
            output_buffer.seek(0)
            
            return FillerResult(
                success=True,
                placeholders_found=self.placeholders_found,
                placeholders_filled=self.placeholders_filled,
                document_bytes=output_buffer.getvalue(),
                unfilled=unfilled,
                errors=errors
            )
            
        except Exception as e:
            logger.error(f"Error filling placeholders: {e}")
            return FillerResult(
                success=False,
                placeholders_found=self.placeholders_found,
                placeholders_filled=self.placeholders_filled,
                unfilled=unfilled,
                errors=[str(e)]
            )
    
    def _find_all_placeholders(self, doc: Document) -> List[str]:
        """Find all unique placeholders in the document."""
        placeholders = set()
        
        # Check paragraphs
        for para in doc.paragraphs:
            matches = self.PLACEHOLDER_PATTERN.findall(para.text)
            placeholders.update(matches)
        
        # Check tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        matches = self.PLACEHOLDER_PATTERN.findall(para.text)
                        placeholders.update(matches)
        
        # Check headers and footers
        for section in doc.sections:
            # Header
            if section.header:
                for para in section.header.paragraphs:
                    matches = self.PLACEHOLDER_PATTERN.findall(para.text)
                    placeholders.update(matches)
            
            # Footer
            if section.footer:
                for para in section.footer.paragraphs:
                    matches = self.PLACEHOLDER_PATTERN.findall(para.text)
                    placeholders.update(matches)
        
        return sorted(list(placeholders))
    
    def _replace_placeholder(self, doc: Document, placeholder: str, content: str) -> bool:
        """Replace a placeholder with content throughout the document."""
        placeholder_text = f"{{{{{placeholder}}}}}"
        replaced = False
        
        # Replace in paragraphs
        for para in doc.paragraphs:
            if placeholder_text in para.text:
                self._replace_in_paragraph(para, placeholder_text, content)
                replaced = True
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if placeholder_text in para.text:
                            self._replace_in_paragraph(para, placeholder_text, content)
                            replaced = True
        
        # Replace in headers and footers
        for section in doc.sections:
            if section.header:
                for para in section.header.paragraphs:
                    if placeholder_text in para.text:
                        self._replace_in_paragraph(para, placeholder_text, content)
                        replaced = True
            
            if section.footer:
                for para in section.footer.paragraphs:
                    if placeholder_text in para.text:
                        self._replace_in_paragraph(para, placeholder_text, content)
                        replaced = True
        
        return replaced
    
    def _replace_in_paragraph(self, para: Paragraph, placeholder: str, content: str) -> None:
        """
        Replace placeholder in a paragraph.
        Handles text that may be split across multiple runs.
        """
        # Try simple run-by-run replacement first
        for run in para.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, content)
                return
        
        # If placeholder spans runs, we need to handle it differently
        # Reconstruct the full text, find and replace, then reset
        full_text = para.text
        if placeholder in full_text:
            new_text = full_text.replace(placeholder, content)
            
            # Clear all runs and set new text in first run
            # This loses some formatting but ensures replacement works
            if para.runs:
                # Keep formatting from first run
                first_run = para.runs[0]
                
                # Clear other runs
                for run in para.runs[1:]:
                    run.text = ""
                
                first_run.text = new_text


def fill_document_placeholders(
    docx_path: str,
    content_mapping: Dict[str, str],
    output_path: Optional[str] = None
) -> FillerResult:
    """
    Convenience function to fill placeholders in a document.
    
    Args:
        docx_path: Path to DOCX template
        content_mapping: Dict of placeholder name -> content
        output_path: Optional output path
        
    Returns:
        FillerResult
    """
    filler = PlaceholderFiller()
    return filler.fill(docx_path, content_mapping, output_path)


if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description="Fill placeholders in DOCX")
    parser.add_argument("--template", required=True, help="Path to DOCX template")
    parser.add_argument("--output", help="Output path for filled document")
    parser.add_argument("--risk", help="Content to fill {{RISK}} placeholder")
    args = parser.parse_args()
    
    content_mapping = {}
    if args.risk:
        content_mapping["RISK"] = args.risk
    
    filler = PlaceholderFiller()
    result = filler.fill(args.template, content_mapping, args.output)
    
    print(f"\nPlaceholders found: {result.placeholders_found}")
    print(f"Placeholders filled: {result.placeholders_filled}")
    if result.unfilled:
        print(f"Unfilled placeholders: {result.unfilled}")
    if result.output_path:
        print(f"Output saved to: {result.output_path}")
