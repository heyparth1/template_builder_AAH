
import os
import sys
from modules.databricks_parser import DatabricksParser
from docx import Document

def test_local_fallback():
    print("Testing local fallback mechanism...")
    
    # Create dummy DOCX
    filename = "test_fallback.docx"
    doc = Document()
    doc.add_paragraph("This is a test document.")
    doc.add_paragraph("It has some content {{PLACEHOLDER}}.")
    doc.save(filename)
    
    abs_path = os.path.abspath(filename)
    print(f"Created dummy file: {abs_path}")
    
    try:
        # Initialize parser (will use mock mode if env vars missing, which we assume or force)
        # We can force mock mode by messing with env vars but let's just rely on it failing to connect if not set,
        # or we explicitly call _mock_response to test that valid part.
        
        parser = DatabricksParser()
        
        # Test 1: Direct _mock_response call
        print("\nTest 1: Calling _mock_response directly...")
        response = parser._mock_response(abs_path)
        
        if response['content'] is None:
            print("FAILED: Content is None")
        else:
            print(f"SUCCESS: Content bytes found ({len(response['content'])} bytes)")
            
        if "test document" in response['parsed']['text']:
            print("SUCCESS: Text extracted correctly")
        else:
            print(f"FAILED: Text not found in: {response['parsed']['text']}")
            
        # Test 2: Full parse_and_classify (which calls _mock_response if connection fails)
        print("\nTest 2: Calling parse_and_classify...")
        # We want to ensure it uses the local path. 
        # In the app, we pass the Volume path OR local path. 
        # If we pass local path here, it should work via fallback.
        
        result = parser.parse_and_classify(abs_path)
        
        if result.document_bytes:
            print(f"SUCCESS: Result has document bytes ({len(result.document_bytes)})")
        else:
            print("FAILED: Result has no document bytes")
            
        if result.classification_result:
            print(f"SUCCESS: Classification run (replacements: {len(result.classification_result.replacements)})")
        else:
            print("FAILED: No classification result")

    except Exception as e:
        print(f"EXCEPTION: {e}")
        import traceback
        traceback.print_exc()
    finally:
        # Cleanup
        if os.path.exists(filename):
            os.remove(filename)
            print(f"\nCleaned up {filename}")

if __name__ == "__main__":
    test_local_fallback()
